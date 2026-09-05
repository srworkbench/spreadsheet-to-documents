"""Generate editable Word documents from a CSV or XLSX file and a text template."""
from __future__ import annotations

import argparse
import csv
import re
import shutil
import tempfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from openpyxl import load_workbook

TOKEN = re.compile(r"\{\{([A-Za-z_][A-Za-z0-9_]*)\}\}")
IDENTIFIER = re.compile(r"[A-Za-z0-9][A-Za-z0-9_-]{0,79}\Z")


def read_records(path: Path, sheet: str | None = None) -> list[dict[str, str]]:
    if path.suffix.lower() == ".csv":
        if sheet:
            raise ValueError("--sheet is only supported for XLSX input")
        with path.open(encoding="utf-8-sig", newline="") as stream:
            rows = list(csv.reader(stream))
    elif path.suffix.lower() == ".xlsx":
        workbook = load_workbook(path, read_only=True, data_only=False)
        try:
            if sheet and sheet not in workbook.sheetnames:
                raise ValueError("Requested worksheet does not exist")
            worksheet = workbook[sheet] if sheet else workbook.active
            rows = []
            for cells in worksheet.iter_rows():
                if any(cell.data_type == "f" for cell in cells):
                    raise ValueError("Formula cells are unsupported; use a values-only workbook")
                rows.append([cell.value for cell in cells])
        finally:
            workbook.close()
    else:
        raise ValueError("Input must be a .csv or .xlsx file")
    if not rows:
        raise ValueError("Input is empty")
    headers = [str(value).strip() if value is not None else "" for value in rows[0]]
    if not all(headers) or len(set(headers)) != len(headers):
        raise ValueError("Headers must be nonempty and unique")
    if "record_id" not in headers:
        raise ValueError("Input requires a record_id column")
    records = []
    for number, row in enumerate(rows[1:], 2):
        if not any(value is not None and str(value).strip() for value in row):
            continue
        if len(row) != len(headers):
            raise ValueError(f"Row {number} has the wrong number of columns")
        records.append(dict(zip(headers, [str(v).strip() if v is not None else "" for v in row])))
    if not records:
        raise ValueError("Input contains no records")
    return records


def render_text(template: str, records: list[dict[str, str]]) -> list[tuple[str, str]]:
    if not template.strip():
        raise ValueError("Template is empty")
    fields = set(TOKEN.findall(template))
    remainder = TOKEN.sub("", template)
    if "{{" in remainder or "}}" in remainder:
        raise ValueError("Malformed template placeholder")
    identifiers: set[str] = set()
    rendered = []
    for number, record in enumerate(records, 2):
        identifier = record["record_id"]
        if not IDENTIFIER.fullmatch(identifier):
            raise ValueError(f"Row {number} has an invalid record_id")
        if identifier.casefold() in identifiers:
            raise ValueError(f"Row {number} repeats a record_id")
        identifiers.add(identifier.casefold())
        if any(not record.get(field) for field in fields):
            raise ValueError(f"Row {number} is missing a required template value")
        text = TOKEN.sub(lambda match: record[match.group(1)], template)
        if re.search(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", text):
            raise ValueError(f"Row {number} contains unsupported control characters")
        rendered.append((identifier, text))
    return rendered


def write_document(text: str, path: Path) -> None:
    document = Document()
    for section in document.sections:
        section.top_margin = section.bottom_margin = Inches(0.8)
        section.left_margin = section.right_margin = Inches(0.9)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(10)
    title = document.styles["Title"]
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.font.underline = False
    for border in list(title.element.iter(qn("w:pBdr"))):
        border.getparent().remove(border)
    blocks = re.split(r"\n\s*\n", text.strip())
    document.add_paragraph(blocks[0], "Title")
    for block in blocks[1:]:
        document.add_paragraph(block)
    props = document.core_properties
    props.author = props.last_modified_by = ""
    props.comments = props.subject = props.keywords = props.category = ""
    props.identifier = props.version = props.language = ""
    props.title = blocks[0]
    props.created = props.modified = datetime(2000, 1, 1, tzinfo=timezone.utc)
    for key, relationship in list(document.part.package.rels.items()):
        if relationship.reltype.endswith("/thumbnail"):
            del document.part.package.rels[key]
    document.save(path)


def generate(source: Path, template_path: Path, output: Path, sheet: str | None = None) -> int:
    if output.exists():
        raise ValueError("Output directory already exists; choose a new path")
    rendered = render_text(template_path.read_text(encoding="utf-8"), read_records(source, sheet))
    output.parent.mkdir(parents=True, exist_ok=True)
    staging = Path(tempfile.mkdtemp(prefix=".document-build-", dir=output.parent))
    try:
        for identifier, content in rendered:
            write_document(content, staging / f"{identifier}.docx")
        # Reserve the destination exclusively, then move the completed files into it.
        # If another process has created it, mkdir fails without altering that directory.
        output.mkdir()
        try:
            for path in staging.iterdir():
                path.rename(output / path.name)
        except BaseException:
            shutil.rmtree(output)
            raise
    finally:
        shutil.rmtree(staging, ignore_errors=True)
    return len(rendered)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", type=Path)
    parser.add_argument("template", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--sheet", help="XLSX worksheet name; defaults to the active sheet")
    args = parser.parse_args()
    try:
        count = generate(args.source, args.template, args.output, args.sheet)
    except (ValueError, OSError) as exc:
        parser.exit(2, f"Cannot generate documents: {exc}\n")
    print(f"Created {count} documents")


if __name__ == "__main__":
    main()
